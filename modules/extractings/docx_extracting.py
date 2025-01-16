import sys
sys.path.append('../../../demo-5')

import os
import docx
from docx.oxml.ns import qn
from win32com import client
from langchain_core.documents import Document
from docx.opc.constants import RELATIONSHIP_TYPE

from modules.extractings import BaseExtractor


class DOCXExtractor(BaseExtractor):

    def __init__(self):
        super().__init__('DOCXExtractor')
        self.result = {
            'content': ''
        }

    def convert_doc_to_docx(self, doc_path):
            return None

    def load(self, docx_path):
        docx_path = os.path.abspath(docx_path)
        try:
            doc = docx.Document(docx_path)
            content = ''
            hyperlinks = []

            def extract_hyperlinks(paragraph):
                """
                Extract hyperlinks from a paragraph.
                """
                hyperlinks = []
                for rel in paragraph._element.xpath('.//w:hyperlink'):
                    rid = rel.get(qn('r:id'))
                    if rid:
                        # Access the relationship and extract the URL
                        rels = paragraph.part.rels
                        if rid in rels and rels[rid].reltype == RELATIONSHIP_TYPE.HYPERLINK:
                            hyperlinks.append(rels[rid].target_ref)
                return hyperlinks

            def iter_block_items(parent):
                """
                Iterate over paragraphs and tables in document order.
                """
                for child in parent.element.body:
                    if child.tag == qn('w:tbl'):  # Table element
                        yield 'table', child
                    elif child.tag == qn('w:p'):  # Paragraph element
                        yield 'paragraph', child

            table_index = 0  # Keep track of tables

            for block_type, block in iter_block_items(doc):
                if block_type == 'paragraph':  # Process paragraph
                    paragraph_text = ''.join(node.text for node in block.xpath('.//w:t') if node.text)
                    hyperlinks = extract_hyperlinks(doc.paragraphs[table_index])
                    if paragraph_text.strip():  # Skip empty paragraphs                    
                        content += paragraph_text.strip() + "\n"

                elif block_type == 'table':  # Process table
                    table = doc.tables[table_index]
                    for row in table.rows:
                        row_data = [cell.text.strip() if cell.text else "" for cell in row.cells]                    
                        content += "\t".join(row_data) + "\n"                
                    content += "\n"
                    table_index += 1  # Move to the next table
            
            if not content.strip():
                print(f"No content found in {os.path.basename(docx_path)}. Skipping...")
                return Document(
                page_content="",
                metadata={
                    "source": docx_path,
                    "hyperlinks": hyperlinks,
                    "tables": [],
                    "error": f'No content found in {docx_path}'
                }
            )
            else:
                document = Document(
                    page_content=content,
                    metadata={
                        "source": docx_path,
                        "hyperlinks": hyperlinks,  
                        "tables": [],  
                        "error": ""
                    }
                )

            return document
        
        except Exception as e:
            error_message = str(e)
            print(f"Error processing {os.path.basename(docx_path)}: {error_message}")
            return Document(
                page_content="",
                metadata={
                    "source": docx_path,
                    "hyperlinks": hyperlinks,
                    "tables": [],
                    "error": error_message
                }
            )
    
    def loads(self, docx_dir):
        documents = []
        for docx_file in os.listdir(docx_dir):
            if docx_file.endswith(".doc"):
                docx_path = os.path.join(docx_dir, docx_file)
                docx_path = self.convert_doc_to_docx(docx_path)
                if docx_path:
                    documents.append(self.load(docx_path))
            elif docx_file.endswith(".docx"):
                docx_path = os.path.join(docx_dir, docx_file)
                documents.append(self.load(docx_path))
        return documents


if __name__ == '__main__':
    docx_extractor = DOCXExtractor()
    doc = docx_extractor.load('../../sources/working/input.docx')

    print(doc)


